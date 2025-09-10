import streamlit as st
from docx import Document
import re
import json
from sentence_transformers import SentenceTransformer, util
import pandas as pd
import os  # ← esta línea es la que faltaba

# Lista de campos clave
CAMPO_CLAVES = [
    "Número de expediente", "Número expediente", "Nombre del representante",
    "N.I.F.", "Empresa", "Razón social", "C.I.F.", "D.N.I.", "Don/Doña",
    "Apoderado", "Poderdante", "Declaración responsable", "Expediente",
    "Firma", "Teléfono", "Correo electrónico", "Provincia", "Localidad", "Código postal"
]

# Modelo semántico
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Funciones de utilidad
def cargar_diccionario(ruta="campos.json"):
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return {}

def guardar_diccionario(diccionario, ruta="campos.json"):
    with open(ruta, "w", encoding="utf-8") as f:
        json.dump(diccionario, f, ensure_ascii=False, indent=2)

def extraer_campos_streamlit(doc, campo_claves):
    campos = set()
    parrafos = [p.text.strip() for p in doc.paragraphs if p.text.strip() != '']

    for i, texto in enumerate(parrafos):
        for clave in campo_claves:
            if clave.lower() in texto.lower():
                campos.add(clave)
        if re.search(r'([^\n\r]+?)(\.{5,}|_+|\s{8,})', texto):
            campos.add(re.sub(r'(\.+|_+|\s{8,})', '', texto).strip(": "))
        if texto.endswith(':'):
            campos.add(texto.strip(": "))
        if i + 1 < len(parrafos):
            siguiente = parrafos[i + 1]
            if len(siguiente) < 5:
                campos.add(texto.strip(": "))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texto = cell.text.strip()
                if texto:
                    for clave in campo_claves:
                        if clave.lower() in texto.lower():
                            campos.add(clave)
                    if re.search(r'([^\n\r]+?)(\.{5,}|_+|\s{8,})', texto):
                        campos.add(re.sub(r'(\.+|_+|\s{8,})', '', texto).strip(": "))
                    if texto.endswith(':'):
                        campos.add(texto.strip(": "))

    campos_filtrados = {c for c in campos if 2 < len(c) < 60}
    return sorted(campos_filtrados)

def sugerir_sinonimos(nuevo_campo, diccionario):
    scores = {}
    for clave in diccionario:
        score = util.cos_sim(model.encode(nuevo_campo), model.encode(clave))
        scores[clave] = score.item()
    sugerido = max(scores, key=scores.get)
    if scores[sugerido] > 0.8:
        return sugerido
    return None

def reemplazar_en_parrafo(parrafo, datos_fila):
    texto_original = "".join(run.text for run in parrafo.runs)
    texto_modificado = texto_original
    for campo, valor in datos_fila.items():
        placeholder = f'{{{{{campo}}}}}'
        texto_modificado = texto_modificado.replace(placeholder, str(valor))
    if texto_modificado != texto_original:
        for run in parrafo.runs:
            run.text = ""
        if parrafo.runs:
            parrafo.runs[0].text = texto_modificado
        else:
            parrafo.add_run(texto_modificado)
def reemplazar_en_celda(celda, datos_fila):
    # Reemplazar en párrafos normales
    for parrafo in celda.paragraphs:
        reemplazar_en_parrafo(parrafo, datos_fila)

    # Reemplazar en tablas anidadas (si las hay)
    for tabla in celda.tables:
        for fila in tabla.rows:
            for subcelda in fila.cells:
                reemplazar_en_celda(subcelda, datos_fila)  # Recursivo
    #obtener ruta de descargas
import platform

def obtener_ruta_descargas():
    home = os.path.expanduser("~")
    sistema = platform.system()
    if sistema == "Windows":
        return os.path.join(home, "Downloads")
    elif sistema == "Darwin":  # macOS
        return os.path.join(home, "Downloads")
    else:  # Linux o desconocido
        return os.path.join(home, "Descargas")  # o "Downloads" si el sistema está en inglés



# Interfaz principal
st.set_page_config(page_title="Gestor de Formularios", layout="wide")
st.sidebar.title("📁 Navegación")
opcion = st.sidebar.radio("Selecciona una sección:", [
    "🔍 Detección de campos",
    "🧠 Diccionario de campos",
    "📝 Rellenar plantillas",
    # "🧪 Generar plantilla",
    # "🖱️ Insertar claves manualmente",
    "📋 Pegado manual de claves"
])



# Sección 1: Detección de campos
if opcion == "🔍 Detección de campos":
    st.title("🧠 Detector inteligente de campos en formularios Word")
    st.write("Sube un archivo .docx y detectaremos automáticamente los campos a rellenar. Podrás confirmar, editar y enriquecer el diccionario.")

    archivo = st.file_uploader("📤 Sube tu archivo Word", type=["docx"])

    if archivo:
        doc = Document(archivo)
        campos_detectados = extraer_campos_streamlit(doc, CAMPO_CLAVES)
        diccionario = cargar_diccionario()

        st.subheader("🔍 Revisión de campos detectados:")
        diccionario_actualizado = False

        for campo in campos_detectados:
            sugerido = sugerir_sinonimos(campo, diccionario)
            if sugerido and sugerido != campo:
                st.markdown(f"- **{campo}** (¿Sinónimo de: _{sugerido}_?)")
                if st.checkbox(f"✅ Confirmar como sinónimo de '{sugerido}'", key=f"sin_{campo}"):
                    diccionario[sugerido].append(campo)
                    diccionario_actualizado = True
            else:
                st.markdown(f"- **{campo}** (Campo nuevo)")
                if st.checkbox(f"🆕 Añadir como nuevo campo", key=f"nuevo_{campo}"):
                    diccionario[campo] = []
                    diccionario_actualizado = True

        if diccionario_actualizado:
            guardar_diccionario(diccionario)
            st.success("🎉 Diccionario actualizado con los cambios confirmados.")

# Sección 2: Diccionario de campos
elif opcion == "🧠 Diccionario de campos":
    st.title("🧠 Diccionario de campos clave")
    diccionario = cargar_diccionario()

    if not diccionario:
        st.info("El diccionario está vacío. Se irá llenando automáticamente al analizar documentos.")
    else:
        for campo, sinonimos in diccionario.items():
            with st.expander(f"📌 {campo}"):
                st.write("Sinónimos actuales:", sinonimos)
                nuevo_sinonimo = st.text_input(f"Añadir sinónimo para '{campo}'", key=f"sin_{campo}")
                if nuevo_sinonimo:
                    if nuevo_sinonimo not in sinonimos:
                        diccionario[campo].append(nuevo_sinonimo)
                        guardar_diccionario(diccionario)
                        st.success(f"Sinónimo '{nuevo_sinonimo}' añadido a '{campo}'")

# Sección 3: Rellenar plantillas (pendiente de implementación)
elif opcion == "📝 Rellenar plantillas":
    st.title("📝 Rellenar plantillas Word con datos de Excel")
    st.write("Sube tu archivo Excel con los datos y las plantillas Word con campos tipo {{Nombre}}, {{DNI}}, etc.")

    excel_file = st.file_uploader("📥 Sube tu archivo Excel", type=["xlsx"])
    plantilla_files = st.file_uploader("📄 Sube tus plantillas Word", type=["docx"], accept_multiple_files=True)

    if excel_file and plantilla_files:
        df = pd.read_excel(excel_file)
        output_folder = obtener_ruta_descargas()
        os.makedirs(output_folder, exist_ok=True)

        documentos_generados = []

        for idx, fila in df.iterrows():
            nombre_base = f'documento_{fila["Nombre"]}_{idx+1}'
            for plantilla in plantilla_files:
                plantilla_doc = Document(plantilla)

                # Reemplazo en párrafos
                for parrafo in plantilla_doc.paragraphs:
                    reemplazar_en_parrafo(parrafo, fila)

                # Reemplazo en tablas
                for table in plantilla_doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            reemplazar_en_celda(cell, fila)

                nombre_doc = os.path.join(output_folder, f'{nombre_base}_{plantilla.name}')
                plantilla_doc.save(nombre_doc)
                documentos_generados.append(nombre_doc)
                st.success(f"✅ Generado: {nombre_doc}")

        # Crear ZIP y mostrar botón de descarga
        import zipfile
        zip_path = "documentos_generados.zip"
        with zipfile.ZipFile(zip_path, "w") as zipf:
            for doc_path in documentos_generados:
                zipf.write(doc_path, arcname=os.path.basename(doc_path))

        with open(zip_path, "rb") as f:
            st.download_button(
                label="⬇️ Descargar todos los documentos en ZIP",
                data=f,
                file_name="documentos_generados.zip",
                mime="application/zip"
            )

        st.info(f"📂 Los documentos generados se han guardado en tu carpeta de Descargas: `{output_folder}`")
        st.balloons()

elif opcion == "🧪 Generar plantilla":
    st.title("🧪 Generar plantilla Word desde documento base")
    st.write("Sube un documento Word con campos vacíos (líneas, puntos, etc.) y detectaremos automáticamente los campos para convertirlos en claves rellenables como {{Nombre}}, {{DNI}}, etc.")

    archivo_base = st.file_uploader("📤 Sube tu documento base (.docx)", type=["docx"])
    if archivo_base:
        doc = Document(archivo_base)
        campos_detectados = extraer_campos_streamlit(doc, CAMPO_CLAVES)

        if campos_detectados:
            st.subheader("✏️ Edita los nombres de los campos detectados:")
            campos_personalizados = {}

            for campo in campos_detectados:
                clave_editada = st.text_input(f"Campo detectado: '{campo}' → Clave para plantilla:", value=campo, key=f"edit_{campo}")
                if clave_editada.strip():
                    campos_personalizados[campo] = clave_editada.strip()

            nombre_plantilla = st.text_input("📝 Nombre para la plantilla (sin extensión)", value="plantilla_generada")

            if st.button("📄 Generar plantilla"):
                # aquí irá la función que genera el documento usando campos_personalizados
    
            
                def generar_plantilla_desde_documento(doc, campos_personalizados, ruta_guardar):
                    for parrafo in doc.paragraphs:
                        texto_original = "".join(run.text for run in parrafo.runs)
                        texto_modificado = texto_original
                        for campo_original, clave_personalizada in campos_personalizados.items():
                            if campo_original in texto_modificado:
                                texto_modificado = texto_modificado.replace(campo_original, f'{campo_original}: {{{{{clave_personalizada}}}}}')
                        if texto_modificado != texto_original:
                            for run in parrafo.runs:
                                run.text = ""
                            if parrafo.runs:
                                parrafo.runs[0].text = texto_modificado
                            else:
                                parrafo.add_run(texto_modificado)

                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    texto_original = "".join(run.text for run in p.runs)
                                    texto_modificado = texto_original
                                    for campo_original, clave_personalizada in campos_personalizados.items():
                                        if campo_original in texto_modificado:
                                            texto_modificado = texto_modificado.replace(campo_original, f'{campo_original}: {{{{{clave_personalizada}}}}}')
                                    if texto_modificado != texto_original:
                                        for run in p.runs:
                                            run.text = ""
                                        if p.runs:
                                            p.runs[0].text = texto_modificado
                                        else:
                                            p.add_run(texto_modificado)

                    doc.save(ruta_guardar)


                carpeta_plantillas = os.path.join(os.path.expanduser("~"), "Downloads", "plantillas_definidas")
                os.makedirs(carpeta_plantillas, exist_ok=True)
                ruta_guardar = os.path.join(carpeta_plantillas, f"{nombre_plantilla}.docx")

                generar_plantilla_desde_documento(doc, campos_personalizados, ruta_guardar)
                st.success(f"🎉 Plantilla generada y guardada en: `{ruta_guardar}`")
                st.balloons()

elif opcion == "🖱️ Insertar claves manualmente":
    st.title("🖱️ Inserta claves manualmente en tu documento Word")
    st.write("Navega por el contenido del documento y selecciona dónde insertar claves como {{Nombre}}, {{DNI}}, etc.")

    archivo_manual = st.file_uploader("📤 Sube tu documento Word (.docx)", type=["docx"])
    lista_de_claves = ["Nombre", "Apellido1", "DNI", "Empresa", "CIF", "Dirección", "Teléfono", "Email"]  # Puedes ampliar esta lista

    if archivo_manual:
        doc = Document(archivo_manual)
        cambios = []

        st.subheader("📄 Párrafos del documento")
        for i, parrafo in enumerate(doc.paragraphs):
            st.markdown(f"**Párrafo {i+1}:** {parrafo.text}")
            if parrafo.text.strip():
                insertar = st.checkbox(f"📌 Insertar clave en este párrafo", key=f"p_check_{i}")
                if insertar:
                    posicion = st.radio("Posición:", ["Inicio", "Final", "Reemplazar"], key=f"p_pos_{i}")
                    clave = st.selectbox("Selecciona clave:", lista_de_claves, key=f"p_clave_{i}")
                    cambios.append(("parrafo", i, posicion, clave))

        st.subheader("📊 Celdas en tablas")
        for t_idx, table in enumerate(doc.tables):
            for r_idx, row in enumerate(table.rows):
                for c_idx, cell in enumerate(row.cells):
                    texto = cell.text.strip()
                    if texto:
                        st.markdown(f"**Tabla {t_idx+1}, Fila {r_idx+1}, Celda {c_idx+1}:** {texto}")
                        insertar = st.checkbox(f"📌 Insertar clave en esta celda", key=f"c_check_{t_idx}_{r_idx}_{c_idx}")
                        if insertar:
                            posicion = st.radio("Posición:", ["Inicio", "Final", "Reemplazar"], key=f"c_pos_{t_idx}_{r_idx}_{c_idx}")
                            clave = st.selectbox("Selecciona clave:", lista_de_claves, key=f"c_clave_{t_idx}_{r_idx}_{c_idx}")
                            cambios.append(("celda", (t_idx, r_idx, c_idx), posicion, clave))

        if st.button("📄 Generar documento con claves"):
            for tipo, ubicacion, posicion, clave in cambios:
                if tipo == "parrafo":
                    parrafo = doc.paragraphs[ubicacion]
                    texto = parrafo.text.strip()
                    if posicion == "Inicio":
                        nuevo = f'{{{{{clave}}}}} {texto}'
                    elif posicion == "Final":
                        nuevo = f'{texto} {{{{{clave}}}}}'
                    else:
                        nuevo = f'{{{{{clave}}}}}'
                    parrafo.clear()
                    parrafo.add_run(nuevo)

                elif tipo == "celda":
                    t_idx, r_idx, c_idx = ubicacion
                    cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
                    for p in cell.paragraphs:
                        texto = p.text.strip()
                        if posicion == "Inicio":
                            nuevo = f'{{{{{clave}}}}} {texto}'
                        elif posicion == "Final":
                            nuevo = f'{texto} {{{{{clave}}}}}'
                        else:
                            nuevo = f'{{{{{clave}}}}}'
                        p.clear()
                        p.add_run(nuevo)

            carpeta_destino = os.path.join(os.path.expanduser("~"), "Downloads", "documentos_editados")
            os.makedirs(carpeta_destino, exist_ok=True)
            ruta_guardar = os.path.join(carpeta_destino, "documento_con_claves.docx")
            doc.save(ruta_guardar)
            st.success(f"🎉 Documento generado con claves insertadas en: `{ruta_guardar}`")
    st.sidebar.title("📁 Navegación")

elif opcion == "📋 Pegado manual de claves":
    st.title("📋 Pegado manual de claves")
    st.write("""
    Sube tu archivo Excel con los datos. Las columnas se convertirán en claves rellenables como `{{Nombre}}`, `{{DNI}}`, etc.
    Abre tu documento Word en paralelo y copia directamente las claves desde aquí para pegarlas donde quieras.
    """)

    excel_file = st.file_uploader("📥 Sube tu archivo Excel", type=["xlsx"])

    if excel_file:
        df = pd.read_excel(excel_file)
        claves_disponibles = df.columns.tolist()

        st.subheader("🧠 Claves disponibles para copiar")

        for clave in claves_disponibles:
            st.text_input(
                label=f"🔹 {clave}",
                value=f"{{{{{clave}}}}}",
                key=f"clave_{clave}"
            )
        st.caption("📌 Usa Ctrl+C o clic derecho para copiar la clave que necesites")
    else:
        st.info("Por favor, sube un archivo Excel para mostrar las claves disponibles.")




    
